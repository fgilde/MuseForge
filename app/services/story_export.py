"""Export a story to Markdown, plain text, DOCX or PDF.

Either as one document or one file per chapter, in which case the files are
returned as a ZIP so a browser download stays a single click.

DOCX uses python-docx and PDF uses reportlab. Both are imported lazily with
an actionable message: a missing optional package must disable one format,
not the whole export.
"""

import io
import os
import re
import zipfile

# Rough page estimate, matching story_pipeline.WORDS_PER_PAGE.
WORDS_PER_PAGE = 275

FORMATS = ("md", "txt", "docx", "pdf")
_MIME = {
    "md": "text/markdown",
    "txt": "text/plain",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "zip": "application/zip",
}


class ExportError(RuntimeError):
    """A format could not be produced — message is meant for the user."""


def mime_for(fmt: str) -> str:
    return _MIME.get(fmt, "application/octet-stream")


def safe_filename(text: str, fallback: str = "story") -> str:
    """A filesystem-safe basename. Windows forbids more than POSIX does, and
    a story title is user input, so be strict rather than clever."""
    cleaned = re.sub(r"[^\w\s.-]", "", (text or "").strip(), flags=re.UNICODE)
    cleaned = re.sub(r"\s+", "_", cleaned).strip("._")
    return (cleaned or fallback)[:80]


def chapter_view(chapter: dict, lang: str | None, original_lang: str) -> tuple[str, str]:
    """(title, text) of a chapter in the requested language.

    Falls back to the original when a translation is missing, so an export
    never silently produces empty chapters.
    """
    if lang and lang != original_lang:
        tr = (chapter.get("translations") or {}).get(lang) or {}
        if tr.get("text"):
            return (tr.get("title") or chapter.get("title") or "Untitled",
                    tr.get("text") or "")
    return (chapter.get("title") or "Untitled", chapter.get("text") or "")


def _chapters_for_export(story: dict, lang: str | None) -> list[tuple[int, str, str]]:
    original = ((story.get("params") or {}).get("language")) or "en"
    out = []
    for chapter in (story.get("chapters") or []):
        title, text = chapter_view(chapter, lang, original)
        if not (text or "").strip():
            continue
        out.append((int(chapter.get("index", len(out))), title, text))
    return out


# ── Renderers. Each returns bytes. ────────────────────────────────────


def _render_md(title: str, chapters: list[tuple[int, str, str]], story: dict) -> bytes:
    parts = [f"# {title}\n"]
    logline = ((story.get("outline") or {}).get("logline") or "").strip()
    if logline:
        parts.append(f"*{logline}*\n")
    for _idx, ctitle, text in chapters:
        parts.append(f"\n## {ctitle}\n\n{text.strip()}\n")
    return "".join(parts).encode("utf-8")


def _render_txt(title: str, chapters: list[tuple[int, str, str]], story: dict) -> bytes:
    parts = [title, "=" * len(title), ""]
    for _idx, ctitle, text in chapters:
        parts += [ctitle, "-" * len(ctitle), "", text.strip(), ""]
    return "\n".join(parts).encode("utf-8")


def _render_docx(title: str, chapters: list[tuple[int, str, str]], story: dict) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise ExportError(
            "DOCX export needs the python-docx package (pip install python-docx==1.1.2)."
        ) from e

    doc = Document()
    doc.add_heading(title, level=0)
    logline = ((story.get("outline") or {}).get("logline") or "").strip()
    if logline:
        para = doc.add_paragraph(logline)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].italic = True

    for position, (_idx, ctitle, text) in enumerate(chapters):
        # Each chapter starts on its own page — that is what makes the file
        # usable as a manuscript rather than one long scroll.
        if position > 0:
            doc.add_page_break()
        doc.add_heading(ctitle, level=1)
        for block in re.split(r"\n\s*\n", text.strip()):
            if block.strip():
                doc.add_paragraph(block.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_pdf(title: str, chapters: list[tuple[int, str, str]], story: dict) -> bytes:
    try:
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A5
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            PageBreak, Paragraph, SimpleDocTemplate, Spacer,
        )
    except ImportError as e:
        raise ExportError(
            "PDF export needs the reportlab package (pip install reportlab==4.2.5)."
        ) from e
    from xml.sax.saxutils import escape

    buffer = io.BytesIO()
    # A5 with generous margins reads like a book rather than a report.
    doc = SimpleDocTemplate(
        buffer, pagesize=A5, title=title,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body", parent=styles["BodyText"], fontSize=10.5, leading=15.5,
        alignment=TA_JUSTIFY, firstLineIndent=5 * mm,
        spaceAfter=0,
    )
    first_body = ParagraphStyle("FirstBody", parent=body, firstLineIndent=0)
    heading = ParagraphStyle(
        "ChapterHeading", parent=styles["Heading1"], fontSize=15, leading=20,
        spaceAfter=8 * mm, alignment=TA_CENTER,
    )

    flow = [
        Paragraph(escape(title), styles["Title"]),
    ]
    logline = ((story.get("outline") or {}).get("logline") or "").strip()
    if logline:
        flow.append(Paragraph(
            escape(logline),
            ParagraphStyle("Logline", parent=styles["Italic"], alignment=TA_CENTER),
        ))
    flow.append(PageBreak())

    for position, (_idx, ctitle, text) in enumerate(chapters):
        if position > 0:
            flow.append(PageBreak())
        flow.append(Paragraph(escape(ctitle), heading))
        blocks = [b.strip() for b in re.split(r"\n\s*\n", text.strip()) if b.strip()]
        for i, block in enumerate(blocks):
            # Typographic convention: the first paragraph of a chapter is
            # not indented, the following ones are.
            flow.append(Paragraph(escape(block).replace("\n", " "),
                                  first_body if i == 0 else body))
        flow.append(Spacer(1, 2 * mm))

    doc.build(flow)
    return buffer.getvalue()


_RENDERERS = {"md": _render_md, "txt": _render_txt,
              "docx": _render_docx, "pdf": _render_pdf}


# ── Public API ────────────────────────────────────────────────────────


def render_story(story: dict, fmt: str, lang: str | None = None) -> tuple[bytes, str]:
    """The whole story as one document. Returns (data, filename)."""
    fmt = (fmt or "md").lower()
    if fmt not in _RENDERERS:
        raise ExportError(f"Unsupported format: {fmt}. Use one of {', '.join(FORMATS)}.")
    chapters = _chapters_for_export(story, lang)
    if not chapters:
        raise ExportError("This story has no written chapters yet.")
    title = story.get("title") or "Untitled story"
    data = _RENDERERS[fmt](title, chapters, story)
    suffix = f"_{lang}" if lang and lang != ((story.get("params") or {}).get("language") or "en") else ""
    return data, f"{safe_filename(title)}{suffix}.{fmt}"


def render_chapter(story: dict, index: int, fmt: str,
                   lang: str | None = None) -> tuple[bytes, str]:
    """A single chapter as one document. Returns (data, filename)."""
    fmt = (fmt or "md").lower()
    if fmt not in _RENDERERS:
        raise ExportError(f"Unsupported format: {fmt}. Use one of {', '.join(FORMATS)}.")
    chapters = story.get("chapters") or []
    match = next((c for c in chapters if int(c.get("index", -1)) == int(index)), None)
    if match is None:
        raise ExportError(f"Chapter {index} does not exist.")
    original = ((story.get("params") or {}).get("language")) or "en"
    title, text = chapter_view(match, lang, original)
    if not (text or "").strip():
        raise ExportError(f"Chapter {index} has no text yet.")
    # A single chapter is its own little document: the chapter title is the
    # heading, and there is no logline to repeat.
    data = _RENDERERS[fmt](title, [(int(index), title, text)], {"outline": {}})
    stem = f"{safe_filename(story.get('title') or 'story')}_{int(index) + 1:02d}_{safe_filename(title, 'chapter')}"
    suffix = f"_{lang}" if lang and lang != original else ""
    return data, f"{stem}{suffix}.{fmt}"


def render_chapters_zip(story: dict, fmt: str,
                        lang: str | None = None) -> tuple[bytes, str]:
    """One file per chapter, bundled into a ZIP. Returns (data, filename)."""
    fmt = (fmt or "md").lower()
    if fmt not in _RENDERERS:
        raise ExportError(f"Unsupported format: {fmt}. Use one of {', '.join(FORMATS)}.")
    chapters = _chapters_for_export(story, lang)
    if not chapters:
        raise ExportError("This story has no written chapters yet.")

    title = story.get("title") or "Untitled story"
    original = ((story.get("params") or {}).get("language")) or "en"
    buffer = io.BytesIO()
    # ZIP_DEFLATED because prose compresses well; the docx/pdf members are
    # already compressed but the overhead is negligible.
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for position, (idx, ctitle, text) in enumerate(chapters):
            data = _RENDERERS[fmt](ctitle, [(idx, ctitle, text)], {"outline": {}})
            # Position-based prefix, not the chapter index: the archive
            # should list in reading order even if indices have gaps.
            name = f"{position + 1:02d}_{safe_filename(ctitle, 'chapter')}.{fmt}"
            archive.writestr(name, data)
    suffix = f"_{lang}" if lang and lang != original else ""
    return buffer.getvalue(), f"{safe_filename(title)}{suffix}_chapters_{fmt}.zip"


def available_formats() -> dict:
    """Which formats this install can actually produce.

    The UI uses this to grey out a format instead of offering one that
    fails on click.
    """
    status = {}
    for fmt in FORMATS:
        if fmt in ("md", "txt"):
            status[fmt] = True
            continue
        try:
            if fmt == "docx":
                import docx  # noqa: F401
            else:
                import reportlab  # noqa: F401
            status[fmt] = True
        except ImportError:
            status[fmt] = False
    return status


if __name__ == "__main__":
    story = {
        "title": "Die Stille danach",
        "params": {"language": "de"},
        "outline": {"logline": "Eine Stadt vergisst ihren Namen."},
        "chapters": [
            {"index": 0, "title": "Ankunft", "text": "Erster Absatz.\n\nZweiter Absatz.",
             "translations": {"en": {"title": "Arrival", "text": "First paragraph.\n\nSecond."}}},
            {"index": 1, "title": "Abschied", "text": "Nur ein Absatz."},
            {"index": 2, "title": "Leer", "text": "   "},
        ],
    }

    # -- language view falls back to the original when untranslated
    assert chapter_view(story["chapters"][0], "en", "de")[0] == "Arrival"
    assert chapter_view(story["chapters"][1], "en", "de")[0] == "Abschied", "must fall back"
    assert chapter_view(story["chapters"][0], None, "de")[0] == "Ankunft"

    # -- empty chapters are skipped, so an export never has blank sections
    assert len(_chapters_for_export(story, None)) == 2

    # -- filenames stay safe and carry the language only when translated
    md, name = render_story(story, "md")
    assert name == "Die_Stille_danach.md", name
    assert b"# Die Stille danach" in md and b"## Ankunft" in md
    _en, en_name = render_story(story, "md", lang="en")
    assert en_name.endswith("_en.md"), en_name

    txt, _ = render_story(story, "txt")
    assert b"Ankunft" in txt and b"Nur ein Absatz." in txt

    # -- per-chapter and zip
    _data, cname = render_chapter(story, 1, "md")
    assert cname == "Die_Stille_danach_02_Abschied.md", cname
    zdata, zname = render_chapters_zip(story, "md")
    assert zname.endswith("_chapters_md.zip"), zname
    with zipfile.ZipFile(io.BytesIO(zdata)) as z:
        names = z.namelist()
    assert names == ["01_Ankunft.md", "02_Abschied.md"], names

    # -- an empty format defaults to markdown (matches the signature default)
    _d, default_name = render_story(story, "")
    assert default_name.endswith(".md"), default_name

    # -- error cases are explicit, never silent
    for bad in ("epub", "html", "doc"):
        try:
            render_story(story, bad)
            raise AssertionError(f"{bad!r} should be rejected")
        except ExportError:
            pass
    try:
        render_chapter(story, 99, "md")
        raise AssertionError("missing chapter should raise")
    except ExportError:
        pass
    try:
        render_chapter(story, 2, "md")  # whitespace-only chapter
        raise AssertionError("empty chapter should raise")
    except ExportError:
        pass

    # -- unsafe titles cannot escape the filename
    nasty = dict(story, title="../../etc/pa$$wd: <hack>")
    _d, safe = render_story(nasty, "md")
    assert "/" not in safe and "\\" not in safe and ":" not in safe, safe

    # -- optional formats: produce real bytes when installed, a clear error if not
    formats = available_formats()
    assert formats["md"] and formats["txt"]
    for fmt in ("docx", "pdf"):
        try:
            data, name = render_story(story, fmt)
            assert len(data) > 500 and name.endswith(f".{fmt}"), (fmt, len(data))
            print(f"  {fmt}: {len(data)} bytes")
        except ExportError as e:
            assert not formats[fmt], f"{fmt} reported available but failed: {e}"
            print(f"  {fmt}: not installed — {e}")

    print("story_export self-check: OK")
